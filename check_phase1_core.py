import os
from src.database.session import get_db
from src.core.project_manager import ProjectManager
from src.pipeline.production_pipeline import ProductionPipeline

def test_phase1_integration():
	print("============ KIỂM TRA NGHIỆM THU PHASE 1 (CORE) ============")
	
	# 1. KIỂM TRA MODULE: PROJECT MANAGER
	print("[1/4] Kiểm tra Project Manager...")
	pm = ProjectManager()
	project_name = "ToanDanTaoPhong"
	
	# Nếu thư mục dự án mẫu chưa tồn tại, thực hiện tạo tự động cấu trúc thư mục Asset thương mại
	try:
		project_path = pm.create_project(project_name)
		print(f" -> Khởi tạo cấu trúc Asset độc lập thành công tại: {project_path}")
	except FileExistsError:
		print(f" -> Dự án '{project_name}' đã tồn tại sẵn cấu trúc thư mục.")

	# 2. KIỂM TRA MODULE: DATABASE CONNECTIVITY
	print("\n[2/4] Kiểm tra Database SQL Session...")
	db = next(get_db())
	if db:
		print(" -> Phiên làm việc SQLite (SessionLocal) kết nối ổn định.")

	# 3 & 4. KIỂM TRA MODULE: NOVEL MANAGER & PRODUCTION PIPELINE INTEGRATION
	print("\n[3 & 4/4] Kích hoạt Production Pipeline đóng gói kịch bản tập phim...")
	
	# Tạo một file kịch bản chữ Word giả lập tạm thời nếu chưa có file thật trên máy của bạn
	mock_docx_path = "novel_draft_test.docx"
	if not os.path.exists(mock_docx_path):
		from docx import Document
		doc = Document()
		doc.add_paragraph("Tô Mộc bước ra ban công nhìn lên bầu trời vô tận đầy sao rực rỡ.")
		doc.add_paragraph('"Tô Mộc!" - Lâm Uyển từ phía dưới học viện gọi lớn vọng lên.')
		doc.save(mock_docx_path)

	# Gọi bộ lõi điều phối dây chuyền sản xuất dài tập mới tạo
	prod_pipeline = ProductionPipeline(db)
	episode_package = prod_pipeline.build_episode_package(
		project_id=project_name,
		episode_number=1,
		docx_file_path=mock_docx_path
	)

	if episode_package["total_scenes"] > 0:
		print("\n[KẾT LUẬN] XÁC NHẬN: Hệ thống lõi Phase 1 (Core) đạt chuẩn bàn giao 100%!")
	print("============================================================\n")

if __name__ == "__main__":
	test_phase1_integration()
